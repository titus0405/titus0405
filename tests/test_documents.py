from io import BytesIO

import openpyxl
import pytest
from docx import Document as DocxDocument
from fpdf import FPDF
from telegram import Document

from bot_tele.documents import SUPPORTED_SUFFIXES, extract_text
from bot_tele.errors import UnsupportedFileError


def _doc(name: str) -> Document:
    return Document(file_id="f", file_unique_id="u", file_name=name)


def _xlsx_bytes() -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Sheet1"
    ws.append(["Nama", "Nilai"])
    ws.append(["Budi", 90])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_bytes() -> bytes:
    doc = DocxDocument()
    _ = doc.add_paragraph("Halo dari dokumen.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bytes() -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    _ = pdf.cell(0, 10, "Teks di dalam PDF.")
    return bytes(pdf.output())


def test_xlsx_extraction() -> None:
    text = extract_text(_doc("data.xlsx"), _xlsx_bytes())
    assert "Sheet1" in text
    assert "Budi" in text
    assert "90" in text


def test_docx_extraction() -> None:
    text = extract_text(_doc("doc.docx"), _docx_bytes())
    assert "Halo dari dokumen." in text


def test_pdf_extraction() -> None:
    text = extract_text(_doc("file.pdf"), _pdf_bytes())
    assert "Teks di dalam PDF." in text


def test_unsupported_format_raises() -> None:
    with pytest.raises(UnsupportedFileError):
        _ = extract_text(_doc("image.png"), b"data")


def test_supported_suffixes_constant() -> None:
    assert ".xlsx" in SUPPORTED_SUFFIXES
    assert ".pdf" in SUPPORTED_SUFFIXES
