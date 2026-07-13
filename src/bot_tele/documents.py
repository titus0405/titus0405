"""Extract plain text from uploaded documents (Excel, Word, PDF)."""

from io import BytesIO
from pathlib import Path
from typing import Final

import xlrd
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
from telegram import Document

from .errors import DocumentReadError, UnsupportedFileError

SUPPORTED_SUFFIXES: Final[frozenset[str]] = frozenset(
    {".xlsx", ".xlsm", ".xls", ".docx", ".pdf"}
)


def extract_text(document: Document, raw: bytes) -> str:
    """Return the plain-text content of `document` based on its extension.

    Raises `UnsupportedFileError` for unknown extensions and `DocumentReadError`
    when a known format cannot be parsed.
    """
    name = document.file_name or ""
    suffix = Path(name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        prefix = f"Format '{suffix or 'tanpa ekstensi'}' belum didukung. "
        raise UnsupportedFileError(prefix + "Gunakan salah satu: " + supported + ".")
    try:
        match suffix:
            case ".xlsx" | ".xlsm":
                return _xlsx_to_text(raw)
            case ".xls":
                return _xls_to_text(raw)
            case ".docx":
                return _docx_to_text(raw)
            case ".pdf":
                return _pdf_to_text(raw)
            case _:
                pass
    except Exception as e:
        raise DocumentReadError(f"Gagal membaca isi file: {e}") from e
    return ""


def _xlsx_to_text(raw: bytes) -> str:
    workbook = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"== Lembar: {sheet.title} ==")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell) for cell in row]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _xls_to_text(raw: bytes) -> str:
    book = xlrd.open_workbook(file_contents=raw)
    parts: list[str] = []
    for sheet in book.sheets():
        parts.append(f"== Lembar: {sheet.name} ==")
        for r in range(sheet.nrows):
            cells = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
            if any(cell.strip() for cell in cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _docx_to_text(raw: bytes) -> str:
    doc = DocxDocument(BytesIO(raw))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _pdf_to_text(raw: bytes) -> str:
    reader = PdfReader(BytesIO(raw))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts)
